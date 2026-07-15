#!/usr/bin/env python3
"""Build and validate a transparent manuscript delivery package."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.etree import ElementTree
from xml.sax.saxutils import escape as xml_escape

from workflow_common import atomic_json_write, emit, load_json_yaml, resolve_vault


MANUSCRIPT_CANDIDATES = (
    "07_Manuscript/Manuscript_AuthorVoice.md",
    "07_Manuscript/Manuscript_Humanized.md",
    "07_Manuscript/Manuscript_Polished.md",
    "07_Manuscript/Manuscript_Draft.md",
)
REQUIRED_DOCX_PARTS = {
    "[Content_Types].xml",
    "_rels/.rels",
    "docProps/core.xml",
    "docProps/app.xml",
    "word/document.xml",
    "word/styles.xml",
    "word/_rels/document.xml.rels",
}
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def atomic_text_write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, temp_name = tempfile.mkstemp(prefix=f".{path.name}.", suffix=".tmp", dir=path.parent)
    try:
        with os.fdopen(fd, "w", encoding="utf-8", newline="\n") as handle:
            handle.write(text)
        Path(temp_name).replace(path)
    finally:
        temp_path = Path(temp_name)
        if temp_path.exists():
            temp_path.unlink()


def select_manuscript(vault: Path, explicit: str | None) -> Path:
    if explicit:
        raw = Path(explicit).expanduser()
        path = raw.resolve() if raw.is_absolute() else (vault / raw).resolve()
        if not raw.is_absolute() and not path.is_relative_to(vault):
            raise ValueError(f"manuscript_path_escapes_vault:{explicit}")
        if not path.is_file():
            raise ValueError(f"manuscript_not_found:{path}")
        return path
    for rel in MANUSCRIPT_CANDIDATES:
        path = vault / rel
        if path.is_file():
            return path
    raise ValueError("manuscript_not_found:no_default_candidate")


def plain_inline(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\(([^)]+)\)", lambda match: f"{match.group(1) or 'Figure'} ({match.group(2)})", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", lambda match: f"{match.group(1)} ({match.group(2)})", text)
    text = re.sub(r"(`{1,3}|\*\*|__|~~)", "", text)
    return text.replace("\t", "    ").strip()


def strip_frontmatter(text: str) -> str:
    """Remove one leading YAML frontmatter block from rendered delivery text."""
    if not text.startswith("---\n"):
        return text
    end = text.find("\n---\n", 4)
    return text[end + 5 :] if end >= 0 else text


def markdown_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    in_code = False
    code_lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                blocks.append(("Code", "\n".join(code_lines)))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            blocks.append(("Title" if not blocks and level == 1 else f"Heading{level}", plain_inline(heading.group(2))))
            continue
        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet:
            blocks.append(("ListParagraph", f"• {plain_inline(bullet.group(1))}"))
            continue
        numbered = re.match(r"^\s*(\d+)[.)]\s+(.+)$", line)
        if numbered:
            blocks.append(("ListParagraph", f"{numbered.group(1)}. {plain_inline(numbered.group(2))}"))
            continue
        if line.startswith(">"):
            blocks.append(("Quote", plain_inline(line.lstrip("> "))))
            continue
        if not line.strip():
            blocks.append(("Normal", ""))
            continue
        blocks.append(("Normal", plain_inline(line)))
    if code_lines:
        blocks.append(("Code", "\n".join(code_lines)))
    return blocks


def paragraph_xml(style: str, text: str) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{xml_escape(style)}"/></w:pPr>' if style else ""
    if not text:
        return f"<w:p>{style_xml}</w:p>"
    lines = text.splitlines() or [text]
    runs: list[str] = []
    for index, value in enumerate(lines):
        if index:
            runs.append("<w:r><w:br/></w:r>")
        preserve = ' xml:space="preserve"' if value.startswith(" ") or value.endswith(" ") else ""
        runs.append(f"<w:r><w:t{preserve}>{xml_escape(value)}</w:t></w:r>")
    return f"<w:p>{style_xml}{''.join(runs)}</w:p>"


def docx_xml_parts(blocks: list[tuple[str, str]], title: str) -> dict[str, str]:
    paragraphs = "".join(paragraph_xml(style, text) for style, text in blocks)
    document = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:document xmlns:w="{W_NS}"><w:body>{paragraphs}'
        '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/>'
        '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>'
        "</w:body></w:document>"
    )
    styles = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="{W_NS}">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="宋体"/><w:sz w:val="22"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="32"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="28"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="24"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading3"><w:name w:val="heading 3"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="22"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="ListParagraph"><w:name w:val="List Paragraph"/><w:basedOn w:val="Normal"/></w:style>
  <w:style w:type="paragraph" w:styleId="Quote"><w:name w:val="Quote"/><w:basedOn w:val="Normal"/><w:rPr><w:i/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Code"><w:name w:val="Code"/><w:basedOn w:val="Normal"/><w:rPr><w:rFonts w:ascii="Courier New" w:hAnsi="Courier New"/></w:rPr></w:style>
</w:styles>'''
    content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>'''
    root_rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>'''
    document_rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>'''
    safe_title = xml_escape(title or "Manuscript")
    now = utc_now().replace("+00:00", "Z")
    core = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>{safe_title}</dc:title><dc:creator>Manuscript workflow</dc:creator>
  <dcterms:created xsi:type="dcterms:W3CDTF">{now}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{now}</dcterms:modified>
</cp:coreProperties>'''
    app = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"><Application>Manuscript workflow</Application></Properties>'''
    return {
        "[Content_Types].xml": content_types,
        "_rels/.rels": root_rels,
        "docProps/core.xml": core,
        "docProps/app.xml": app,
        "word/document.xml": document,
        "word/styles.xml": styles,
        "word/_rels/document.xml.rels": document_rels,
    }


def build_docx_minimal(path: Path, markdown: str, title: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, temp_name = tempfile.mkstemp(prefix=f".{path.name}.", suffix=".tmp", dir=path.parent)
    os.close(fd)
    try:
        parts = docx_xml_parts(markdown_blocks(markdown), title)
        with zipfile.ZipFile(temp_name, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for name, content in parts.items():
                archive.writestr(name, content.encode("utf-8"))
        Path(temp_name).replace(path)
    finally:
        temp_path = Path(temp_name)
        if temp_path.exists():
            temp_path.unlink()


def choose_style(document: Any, *names: str) -> str | None:
    available = {style.name for style in document.styles}
    return next((name for name in names if name in available), None)


def clear_document_body(document: Any) -> None:
    body = document._element.body
    for child in list(body):
        if not child.tag.endswith("sectPr"):
            body.remove(child)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines) or "|" not in lines[start]:
        return None
    separator = lines[start + 1].strip()
    if not re.fullmatch(r"\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?", separator):
        return None
    rows = []
    index = start
    while index < len(lines) and "|" in lines[index] and lines[index].strip():
        if index != start + 1:
            rows.append([plain_inline(cell.strip()) for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def build_docx_rich(path: Path, markdown: str, title: str, template: Path | None, source_dir: Path) -> dict[str, Any]:
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt  # type: ignore

    document = Document(str(template)) if template else Document()
    if template:
        clear_document_body(document)
    normal_style = choose_style(document, "论文正文", "Normal")
    title_style = choose_style(document, "Title", "标题")
    heading_styles = {
        1: choose_style(document, "Heading 1", "标题 1"),
        2: choose_style(document, "Heading 2", "标题 2"),
        3: choose_style(document, "Heading 3", "标题 3"),
    }
    list_style = choose_style(document, "List Paragraph", "论文正文", "Normal")
    quote_style = choose_style(document, "Quote", "论文正文", "Normal")
    figure_style = choose_style(document, "图片", "Normal")
    caption_style = choose_style(document, "图例", "Caption", "Normal")
    table_text_style = choose_style(document, "表格内容", "Normal")

    lines = markdown.splitlines()
    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index].rstrip()
        table_data = parse_markdown_table(lines, index)
        if table_data:
            rows, next_index = table_data
            width = max((len(row) for row in rows), default=0)
            if rows and width:
                table = document.add_table(rows=len(rows), cols=width)
                preferred_table_style = choose_style(document, "样式1", "Table Grid")
                if preferred_table_style:
                    table.style = preferred_table_style
                for r_index, row in enumerate(rows):
                    for c_index in range(width):
                        cell = table.cell(r_index, c_index)
                        cell.text = row[c_index] if c_index < len(row) else ""
                        for paragraph in cell.paragraphs:
                            if table_text_style:
                                paragraph.style = table_text_style
                index = next_index
                continue
        image = re.fullmatch(r"\s*!\[([^]]*)\]\(([^)]+)\)\s*", line)
        if image:
            alt, target = image.groups()
            image_path = Path(target).expanduser()
            if not image_path.is_absolute():
                image_path = (source_dir / image_path).resolve()
            paragraph = document.add_paragraph(style=figure_style)
            if image_path.is_file():
                paragraph.add_run().add_picture(str(image_path), width=Inches(6.2))
            else:
                paragraph.add_run(f"[Missing figure: {target}]")
            if alt:
                document.add_paragraph(plain_inline(alt), style=caption_style)
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            text = plain_inline(heading.group(2))
            style = title_style if first_heading else heading_styles.get(level)
            document.add_paragraph(text, style=style)
            first_heading = False
        elif re.match(r"^\s*[-*+]\s+", line):
            document.add_paragraph("• " + plain_inline(re.sub(r"^\s*[-*+]\s+", "", line)), style=list_style)
        elif re.match(r"^\s*\d+[.)]\s+", line):
            document.add_paragraph(plain_inline(line.strip()), style=list_style)
        elif line.startswith(">"):
            document.add_paragraph(plain_inline(line.lstrip("> ")), style=quote_style)
        else:
            document.add_paragraph(plain_inline(line), style=normal_style)
        index += 1
    if not template and normal_style:
        style = document.styles[normal_style]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(f"{{{W_NS}}}eastAsia", "宋体")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return {"mode": "python-docx", "template": str(template) if template else "", "template_used": bool(template)}


def build_docx(path: Path, markdown: str, title: str, template: Path | None, source_dir: Path) -> dict[str, Any]:
    try:
        import docx  # type: ignore  # noqa: F401
    except ImportError:
        if template:
            raise ValueError("word_template_requested_but_python_docx_missing")
        build_docx_minimal(path, markdown, title)
        return {"mode": "stdlib-minimal", "template": "", "template_used": False}
    return build_docx_rich(path, markdown, title, template, source_dir)


def norm_docx_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip().casefold()


def validate_docx(path: Path, expected_title: str = "") -> dict[str, Any]:
    errors: list[str] = []
    paragraph_count = 0
    first_paragraph = ""
    if not path.is_file():
        return {"valid": False, "errors": ["file_missing"], "path": str(path)}
    if path.stat().st_size < 1000:
        errors.append("file_too_small")
    try:
        with zipfile.ZipFile(path, "r") as archive:
            corrupt = archive.testzip()
            if corrupt:
                errors.append(f"corrupt_member:{corrupt}")
            names = set(archive.namelist())
            missing = sorted(REQUIRED_DOCX_PARTS - names)
            if missing:
                errors.extend(f"missing_part:{name}" for name in missing)
            for name in sorted(REQUIRED_DOCX_PARTS & names):
                try:
                    ElementTree.fromstring(archive.read(name))
                except ElementTree.ParseError as exc:
                    errors.append(f"invalid_xml:{name}:{exc}")
            if "word/document.xml" in names:
                root = ElementTree.fromstring(archive.read("word/document.xml"))
                paragraphs = root.findall(f".//{{{W_NS}}}p")
                paragraph_count = len(paragraphs)
                for paragraph in paragraphs:
                    value = "".join(node.text or "" for node in paragraph.findall(f".//{{{W_NS}}}t")).strip()
                    if value:
                        first_paragraph = value
                        break
                if paragraph_count < 1:
                    errors.append("document_has_no_paragraphs")
                if expected_title and norm_docx_text(first_paragraph) != norm_docx_text(expected_title):
                    errors.append("first_paragraph_title_mismatch")
    except (OSError, zipfile.BadZipFile, ElementTree.ParseError) as exc:
        errors.append(f"open_failed:{exc}")
    return {"valid": not errors, "errors": errors, "path": str(path), "paragraph_count": paragraph_count, "first_paragraph": first_paragraph}


LATEX_ESCAPES = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def latex_escape(text: str) -> str:
    return "".join(LATEX_ESCAPES.get(char, char) for char in text)


def latex_inline(text: str) -> str:
    parts = re.split(r"(\$[^$\n]+\$)", text)
    rendered: list[str] = []
    for part in parts:
        if len(part) >= 2 and part.startswith("$") and part.endswith("$"):
            rendered.append(part)
        else:
            value = plain_inline(part)
            rendered.append(latex_escape(value))
    return "".join(rendered)


def markdown_to_latex(markdown: str, source_dir: Path, title: str) -> str:
    body: list[str] = []
    in_code = False
    in_itemize = False
    in_enumerate = False
    title_consumed = False

    def close_lists() -> None:
        nonlocal in_itemize, in_enumerate
        if in_itemize:
            body.append(r"\end{itemize}")
            in_itemize = False
        if in_enumerate:
            body.append(r"\end{enumerate}")
            in_enumerate = False

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.strip().startswith("```"):
            close_lists()
            body.append(r"\end{verbatim}" if in_code else r"\begin{verbatim}")
            in_code = not in_code
            continue
        if in_code:
            body.append(line)
            continue
        image = re.fullmatch(r"\s*!\[([^]]*)\]\(([^)]+)\)\s*", line)
        if image:
            close_lists()
            alt, target = image.groups()
            if re.match(r"https?://", target, flags=re.IGNORECASE):
                body.append(r"\textit{" + latex_escape(alt or "Figure") + ": " + latex_escape(target) + "}")
            else:
                image_path = (source_dir / target).resolve() if not Path(target).is_absolute() else Path(target).resolve()
                tex_path = image_path.as_posix()
                body.extend(
                    [
                        r"\begin{figure}[htbp]",
                        r"\centering",
                        r"\IfFileExists{\detokenize{" + tex_path + r"}}{\includegraphics[width=0.95\linewidth]{\detokenize{" + tex_path + r"}}}{\fbox{Missing figure: " + latex_escape(target) + "}}",
                        r"\caption{" + latex_escape(alt or "Figure") + "}",
                        r"\end{figure}",
                    ]
                )
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            close_lists()
            level = len(heading.group(1))
            value = latex_inline(heading.group(2))
            if level == 1 and not title_consumed:
                title_consumed = True
                continue
            command = {1: "section", 2: "section", 3: "subsection", 4: "subsubsection"}.get(level, "paragraph")
            body.append(f"\\{command}{{{value}}}")
            continue
        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet:
            if in_enumerate:
                body.append(r"\end{enumerate}")
                in_enumerate = False
            if not in_itemize:
                body.append(r"\begin{itemize}")
                in_itemize = True
            body.append(r"\item " + latex_inline(bullet.group(1)))
            continue
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if numbered:
            if in_itemize:
                body.append(r"\end{itemize}")
                in_itemize = False
            if not in_enumerate:
                body.append(r"\begin{enumerate}")
                in_enumerate = True
            body.append(r"\item " + latex_inline(numbered.group(1)))
            continue
        close_lists()
        if not line.strip():
            body.append("")
        elif line.startswith(">"):
            body.append(r"\begin{quote}" + latex_inline(line.lstrip("> ")) + r"\end{quote}")
        else:
            body.append(latex_inline(line) + "\n")
    close_lists()
    if in_code:
        body.append(r"\end{verbatim}")
    safe_title = latex_escape(title or "Manuscript")
    return f'''\\documentclass[11pt]{{article}}
\\usepackage{{iftex}}
\\ifPDFTeX
  \\usepackage[T1]{{fontenc}}
  \\usepackage[utf8]{{inputenc}}
\\else
  \\usepackage{{fontspec}}
\\fi
\\usepackage[margin=2.5cm]{{geometry}}
\\usepackage{{graphicx}}
\\usepackage{{hyperref}}
\\usepackage{{microtype}}
\\title{{{safe_title}}}
\\date{{}}
\\begin{{document}}
\\maketitle
{os.linesep.join(body)}
\\end{{document}}
'''


def validate_latex(path: Path) -> dict[str, Any]:
    errors: list[str] = []
    if not path.is_file():
        return {"valid": False, "errors": ["file_missing"], "path": str(path)}
    text = path.read_text(encoding="utf-8-sig")
    if "\\documentclass" not in text:
        errors.append("documentclass_missing")
    if "\\begin{document}" not in text or "\\end{document}" not in text:
        errors.append("document_boundary_missing")
    if text.index("\\begin{document}") > text.index("\\end{document}") if "\\begin{document}" in text and "\\end{document}" in text else False:
        errors.append("document_boundary_order_invalid")
    return {"valid": not errors, "errors": errors, "path": str(path), "characters": len(text)}


def validate_pdf(path: Path) -> dict[str, Any]:
    errors: list[str] = []
    if not path.is_file():
        return {"valid": False, "errors": ["file_missing"], "path": str(path)}
    size = path.stat().st_size
    with path.open("rb") as handle:
        signature = handle.read(5)
        handle.seek(max(0, size - 4096))
        tail = handle.read()
    if signature != b"%PDF-":
        errors.append("signature_invalid")
    if b"%%EOF" not in tail:
        errors.append("eof_missing")
    if size < 512:
        errors.append("file_too_small")
    return {"valid": not errors, "errors": errors, "path": str(path), "bytes": size}


def command_probe(name: str) -> dict[str, Any]:
    executable = shutil.which(name)
    item: dict[str, Any] = {"available": bool(executable), "path": executable, "version": ""}
    if not executable:
        return item
    try:
        completed = subprocess.run(
            [executable, "--version"], capture_output=True, text=True, timeout=8, check=False, encoding="utf-8", errors="replace"
        )
        output = (completed.stdout or completed.stderr or "").strip().splitlines()
        item["version"] = output[0][:300] if output else ""
    except (OSError, subprocess.TimeoutExpired) as exc:
        item["probe_error"] = str(exc)
    return item


def runtime_probe() -> dict[str, Any]:
    runtime = {name: command_probe(name) for name in ("pandoc", "xelatex", "lualatex", "pdflatex")}
    word_candidates = [
        Path(r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE"),
        Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\WINWORD.EXE"),
    ]
    word = next((path for path in word_candidates if path.is_file()), None)
    runtime["word"] = {"available": bool(word), "path": str(word) if word else "", "version": ""}
    runtime["python_docx"] = {"available": False, "path": "", "version": ""}
    try:
        import docx  # type: ignore
        runtime["python_docx"] = {"available": True, "path": str(Path(docx.__file__).resolve()), "version": getattr(docx, "__version__", "")}
    except ImportError:
        pass
    return runtime


def run_command(command: list[str], cwd: Path, timeout: int) -> dict[str, Any]:
    try:
        completed = subprocess.run(
            command,
            cwd=cwd,
            capture_output=True,
            text=True,
            timeout=timeout,
            check=False,
            encoding="utf-8",
            errors="replace",
        )
        return {
            "command": command,
            "returncode": completed.returncode,
            "stdout_tail": completed.stdout[-4000:],
            "stderr_tail": completed.stderr[-4000:],
        }
    except (OSError, subprocess.TimeoutExpired) as exc:
        return {"command": command, "returncode": None, "runtime_error": str(exc)}


def word_com_export(docx_path: Path, pdf_path: Path, timeout: int) -> dict[str, Any]:
    powershell = shutil.which("powershell.exe") or r"C:\Windows\System32\WindowsPowerShell\v1.0\powershell.exe"
    def quote(value: Path) -> str:
        return str(value.resolve()).replace("'", "''")
    script = (
        "$ErrorActionPreference='Stop'; $word=$null; $doc=$null; "
        "try { $word=New-Object -ComObject Word.Application; $word.Visible=$false; $word.DisplayAlerts=0; "
        f"$doc=$word.Documents.Open('{quote(docx_path)}',$false,$true); "
        f"$doc.ExportAsFixedFormat('{quote(pdf_path)}',17); $doc.Close($false); $doc=$null }} "
        "finally { if($doc -ne $null){$doc.Close($false)}; if($word -ne $null){$word.Quit()}; "
        "[System.GC]::Collect(); [System.GC]::WaitForPendingFinalizers() }"
    )
    result = run_command([powershell, "-NoProfile", "-NonInteractive", "-ExecutionPolicy", "Bypass", "-Command", script], pdf_path.parent, timeout)
    result["strategy"] = "word_com"
    return result


def build_pdf(source: Path, tex_path: Path, docx_path: Path, pdf_path: Path, runtime: dict[str, Any], timeout: int, allow_word_com: bool) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    attempts: list[dict[str, Any]] = []
    has_cjk = bool(re.search(r"[\u3400-\u9fff]", source.read_text(encoding="utf-8-sig")))
    available_engines = [name for name in ("xelatex", "lualatex", "pdflatex") if runtime[name]["available"]]
    if has_cjk:
        available_engines = [name for name in available_engines if name != "pdflatex"]
    if not available_engines:
        if allow_word_com and runtime.get("word", {}).get("available"):
            attempt = word_com_export(docx_path, pdf_path, timeout)
            attempts.append(attempt)
            return validate_pdf(pdf_path), attempts
        reason = "no_unicode_tex_engine" if has_cjk else "no_tex_engine"
        if runtime.get("word", {}).get("available") and not allow_word_com:
            reason += ":word_com_available_but_not_authorized"
        return {"valid": False, "errors": [reason], "path": str(pdf_path)}, attempts
    engine = available_engines[0]
    if runtime["pandoc"]["available"]:
        command = [runtime["pandoc"]["path"], str(source), "-o", str(pdf_path), f"--pdf-engine={runtime[engine]['path']}"]
        attempt = run_command(command, pdf_path.parent, timeout)
        attempt["strategy"] = "pandoc"
        attempts.append(attempt)
        validation = validate_pdf(pdf_path)
        if attempt.get("returncode") == 0 and validation["valid"]:
            return validation, attempts
    command = [
        runtime[engine]["path"],
        "-interaction=nonstopmode",
        "-halt-on-error",
        f"-output-directory={pdf_path.parent}",
        str(tex_path),
    ]
    attempt = run_command(command, pdf_path.parent, timeout)
    attempt["strategy"] = "direct_tex"
    attempts.append(attempt)
    generated = pdf_path.parent / f"{tex_path.stem}.pdf"
    if generated.is_file() and generated.resolve() != pdf_path.resolve():
        generated.replace(pdf_path)
    validation = validate_pdf(pdf_path)
    if not validation["valid"] and allow_word_com and runtime.get("word", {}).get("available"):
        attempt = word_com_export(docx_path, pdf_path, timeout)
        attempts.append(attempt)
        validation = validate_pdf(pdf_path)
    return validation, attempts


def artifact_entry(path: Path, kind: str, validation: dict[str, Any]) -> dict[str, Any]:
    return {
        "type": kind,
        "path": str(path),
        "exists": path.is_file(),
        "bytes": path.stat().st_size if path.is_file() else 0,
        "sha256": sha256_file(path) if path.is_file() else None,
        "valid": bool(validation.get("valid")),
        "validation_errors": validation.get("errors", []),
    }


def extract_title(markdown: str, fallback: str) -> str:
    match = re.search(r"(?m)^#\s+(.+?)\s*$", markdown)
    return plain_inline(match.group(1)) if match else fallback


def build_delivery(args: argparse.Namespace) -> int:
    vault = resolve_vault(args.vault)
    source = select_manuscript(vault, args.manuscript)
    project_manifest = load_json_yaml(vault / "00_Control" / "project_manifest.yaml")
    template_value = args.template if args.template is not None else str(project_manifest.get("word_template_path", "") or "")
    template = None
    if template_value.strip():
        raw_template = Path(template_value).expanduser()
        template = raw_template.resolve() if raw_template.is_absolute() else (vault / raw_template).resolve()
        if not template.is_file() or template.suffix.lower() != ".docx":
            raise ValueError(f"word_template_not_found_or_not_docx:{template}")
    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else vault / "09_Delivery"
    output_dir.mkdir(parents=True, exist_ok=True)
    basename = args.basename or source.stem
    if not re.fullmatch(r"[A-Za-z0-9._-]+", basename):
        raise ValueError("basename_must_use_ascii_letters_numbers_dot_dash_underscore")
    markdown = strip_frontmatter(source.read_text(encoding="utf-8-sig"))
    if len(markdown.strip()) < 500:
        raise ValueError("manuscript_too_short_for_delivery")
    title = extract_title(markdown, source.stem)
    docx_path = output_dir / f"{basename}.docx"
    tex_path = output_dir / f"{basename}.tex"
    pdf_path = output_dir / f"{basename}.pdf"
    manifest_path = output_dir / "delivery_manifest.json"
    manifest_hash_path = output_dir / "delivery_manifest.sha256"

    docx_build = build_docx(docx_path, markdown, title, template, source.parent)
    docx_validation = validate_docx(docx_path, title)
    atomic_text_write(tex_path, markdown_to_latex(markdown, source.parent, title))
    tex_validation = validate_latex(tex_path)
    runtime = runtime_probe()
    pdf_validation, attempts = build_pdf(source, tex_path, docx_path, pdf_path, runtime, args.timeout, args.allow_word_com)

    structural_errors = []
    if not docx_validation["valid"]:
        structural_errors.append("docx_invalid")
    if not tex_validation["valid"]:
        structural_errors.append("latex_invalid")
    if structural_errors:
        status = "invalid"
    elif not pdf_validation["valid"]:
        status = "blocked_by_runtime"
    else:
        status = "ready"

    artifacts = [
        artifact_entry(docx_path, "docx", docx_validation),
        artifact_entry(tex_path, "latex", tex_validation),
        artifact_entry(pdf_path, "pdf", pdf_validation),
    ]
    manifest = {
        "schema_version": 1,
        "workflow": "mining-rock-co2-paper-lifecycle",
        "generated_at": utc_now(),
        "status": status,
        "vault": str(vault),
        "source": {"path": str(source), "bytes": source.stat().st_size, "sha256": sha256_file(source)},
        "title": title,
        "runtime": runtime,
        "docx_build": docx_build,
        "word_template": str(template) if template else "",
        "word_com_authorized": bool(args.allow_word_com),
        "pdf_attempts": attempts,
        "artifacts": artifacts,
        "errors": structural_errors + ([] if pdf_validation["valid"] else [f"pdf:{value}" for value in pdf_validation["errors"]]),
        "manifest_hash_file": str(manifest_hash_path),
    }
    atomic_json_write(manifest_path, manifest)
    manifest_sha256 = sha256_file(manifest_path)
    atomic_text_write(manifest_hash_path, f"{manifest_sha256}  {manifest_path.name}\n")
    emit(
        {
            "status": status,
            "manifest": str(manifest_path),
            "manifest_sha256": manifest_sha256,
            "manifest_hash_file": str(manifest_hash_path),
            "artifacts": artifacts,
            "runtime": runtime,
            "docx_build": docx_build,
            "word_template": str(template) if template else "",
            "pdf_attempts": attempts,
        }
    )
    return 0 if status == "ready" else 1


def validate_delivery(args: argparse.Namespace) -> int:
    if args.delivery_dir:
        delivery_dir = Path(args.delivery_dir).expanduser().resolve()
    elif args.vault:
        delivery_dir = resolve_vault(args.vault) / "09_Delivery"
    else:
        raise ValueError("validate_requires_vault_or_delivery_dir")
    manifest_path = Path(args.manifest).expanduser().resolve() if args.manifest else delivery_dir / "delivery_manifest.json"
    errors: list[str] = []
    if not manifest_path.is_file():
        raise ValueError(f"manifest_not_found:{manifest_path}")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    validations: list[dict[str, Any]] = []
    for item in manifest.get("artifacts", []):
        path = Path(str(item.get("path", ""))).expanduser()
        kind = item.get("type")
        if kind == "docx":
            validation = validate_docx(path)
        elif kind == "latex":
            validation = validate_latex(path)
        elif kind == "pdf":
            validation = validate_pdf(path)
        else:
            validation = {"valid": False, "errors": ["unsupported_artifact_type"], "path": str(path)}
        expected_hash = item.get("sha256")
        actual_hash = sha256_file(path) if path.is_file() else None
        if expected_hash != actual_hash:
            validation.setdefault("errors", []).append("sha256_mismatch")
            validation["valid"] = False
        validation["type"] = kind
        validation["expected_sha256"] = expected_hash
        validation["actual_sha256"] = actual_hash
        validations.append(validation)
        if not validation["valid"]:
            errors.append(f"artifact_invalid:{kind}:{path}")
    hash_path = Path(manifest.get("manifest_hash_file") or delivery_dir / "delivery_manifest.sha256")
    if not hash_path.is_file():
        errors.append("manifest_hash_file_missing")
        expected_manifest_hash = None
    else:
        expected_manifest_hash = hash_path.read_text(encoding="utf-8-sig").strip().split()[0]
        if expected_manifest_hash != sha256_file(manifest_path):
            errors.append("manifest_sha256_mismatch")
    status = "valid" if not errors else "invalid"
    emit(
        {
            "status": status,
            "manifest": str(manifest_path),
            "manifest_expected_sha256": expected_manifest_hash,
            "manifest_actual_sha256": sha256_file(manifest_path),
            "artifacts": validations,
            "errors": errors,
        }
    )
    return 0 if not errors else 1


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="生成并校验 DOCX、LaTeX、PDF 和交付哈希清单。")
    subparsers = parser.add_subparsers(dest="command", required=True)
    build = subparsers.add_parser("build", help="构建交付包")
    build.add_argument("--vault", required=True)
    build.add_argument("--manuscript")
    build.add_argument("--output-dir")
    build.add_argument("--basename")
    build.add_argument("--template", help="可选 Word 模板；未给出时读取 project_manifest.word_template_path")
    build.add_argument("--allow-word-com", action="store_true", help="显式允许在 Windows 上用本机 Word COM 将 DOCX 导出为 PDF")
    build.add_argument("--timeout", type=int, default=120)
    validate = subparsers.add_parser("validate", help="重新校验现有交付包和哈希")
    validate.add_argument("--vault")
    validate.add_argument("--delivery-dir")
    validate.add_argument("--manifest")
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    try:
        if args.command == "build":
            return build_delivery(args)
        return validate_delivery(args)
    except (OSError, ValueError, UnicodeError, json.JSONDecodeError, zipfile.BadZipFile) as exc:
        emit({"status": "error", "command": args.command, "message": str(exc)})
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
